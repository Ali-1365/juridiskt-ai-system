"""
Hjälpfunktioner för det juridiska AI-systemet.
Innehåller funktioner för export, filhantering och annat.
"""

import docx
import streamlit as st
import os
import json
from datetime import datetime
import io

def export_to_docx(text, title="Svar från juridiskt AI-system", filename="juridiskt_svar.docx"):
    """
    Export text to a DOCX file
    
    Args:
        text (str): The text to export
        title (str): The title for the document
        filename (str): The filename to save as
        
    Returns:
        bytes: The document as bytes
    """
    # Skapa ett nytt dokument
    doc = docx.Document()
    
    # Lägg till titel
    doc.add_heading(title, level=1)
    
    # Lägg till datum
    doc.add_paragraph(f"Genererat: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    # Lägg till horisontell linje
    doc.add_paragraph("_" * 50)
    
    # Lägg till innehåll
    # Dela upp texten i stycken
    paragraphs = text.split("\n")
    for para in paragraphs:
        if para.strip():  # Hoppa över tomma rader
            if para.strip().startswith("#"):  # Markdown-rubrik
                # Ta bort # och beräkna rubriknivå
                level = 1
                while para.startswith("#"):
                    level += 1
                    para = para[1:].strip()
                
                # Begränsa level till max 9 (docx-begränsning)
                level = min(level, 9)
                
                # Lägg till rubriken
                doc.add_heading(para, level=level)
            elif para.strip().startswith("*") or para.strip().startswith("-"):  # Bullet points
                doc.add_paragraph(para.strip()[1:].strip(), style="List Bullet")
            elif para.strip().startswith("1.") or para.strip().startswith("1)"):  # Numbered list
                doc.add_paragraph(para.strip()[2:].strip(), style="List Number")
            else:
                doc.add_paragraph(para)
    
    # Spara till en BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def download_docx(text, title="Svar från juridiskt AI-system", filename="juridiskt_svar.docx"):
    """
    Provide a download button for a DOCX document in Streamlit
    
    Args:
        text (str): The text to export
        title (str): The title for the document
        filename (str): The filename to save as
    """
    buffer = export_to_docx(text, title, filename)
    
    st.download_button(
        label="⬇️ Ladda ner som DOCX",
        data=buffer,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def save_to_exports(content, filename, file_type="txt"):
    """
    Save content to the exports directory
    
    Args:
        content (str): The content to save
        filename (str): The filename to save as
        file_type (str): The file type (txt, md, json)
        
    Returns:
        str: The path to the saved file
    """
    # Skapa exports-mappen om den inte finns
    os.makedirs("exports", exist_ok=True)
    
    # Skapa fullständig sökväg
    filepath = os.path.join("exports", filename)
    
    # Spara innehållet
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    
    return filepath

def save_interaction_history(username, prompt, response, model="gpt-4o"):
    """
    Save an interaction to the history file
    
    Args:
        username (str): The username
        prompt (str): The user's prompt
        response (str): The AI's response
        model (str): The model used
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Skapa mappen om den inte finns
        os.makedirs("logs", exist_ok=True)
        
        # Skapa filnamn med datum
        date_str = datetime.now().strftime("%Y-%m-%d")
        filename = f"logs/history_{date_str}.jsonl"
        
        # Skapa historik-entry
        entry = {
            "username": username,
            "prompt": prompt,
            "response": response,
            "model": model,
            "timestamp": datetime.now().isoformat()
        }
        
        # Lägg till i filen (append-mode)
        with open(filename, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
        
        return True
    except Exception as e:
        print(f"Fel vid sparande av historik: {str(e)}")
        return False

def create_export_file(data, format="json", prefix="export"):
    """
    Create an export file with given data
    
    Args:
        data: The data to export (dict, list, or string)
        format (str): The format to export as (json, txt, md)
        prefix (str): The prefix for the filename
        
    Returns:
        str: The path to the exported file
    """
    # Skapa mappen om den inte finns
    os.makedirs("exports", exist_ok=True)
    
    # Skapa filnamn med datum och tid
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{prefix}_{timestamp}.{format}"
    filepath = os.path.join("exports", filename)
    
    # Spara data baserat på format
    if format == "json":
        with open(filepath, "w", encoding="utf-8") as f:
            if isinstance(data, (dict, list)):
                json.dump(data, f, ensure_ascii=False, indent=2)
            else:
                f.write(str(data))
    else:
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(str(data))
    
    return filepath

def format_timestamp(timestamp_str, format="long"):
    """
    Format a timestamp string to a more readable format
    
    Args:
        timestamp_str (str): ISO timestamp string
        format (str): Format type (long, short, date, time)
        
    Returns:
        str: Formatted timestamp
    """
    try:
        dt = datetime.fromisoformat(timestamp_str)
        
        if format == "short":
            return dt.strftime("%Y-%m-%d %H:%M")
        elif format == "date":
            return dt.strftime("%Y-%m-%d")
        elif format == "time":
            return dt.strftime("%H:%M:%S")
        else:  # long
            return dt.strftime("%Y-%m-%d %H:%M:%S")
    except:
        return timestamp_str