"""
DOCX-mallgenerator – skapar juridiska dokument som DOCX.
Använder python-docx. Format: Överklagande, JO-anmälan m.m.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_docx(data: dict, template_type: str = "överklagande", filename="utkast.docx"):
    doc = Document()
    doc.add_heading(template_type.capitalize(), level=1)

    intro = f"Till: {data.get('mottagare', 'Förvaltningsrätten')}"
    doc.add_paragraph(intro)

    p1 = doc.add_paragraph()
    p1.add_run(f"Jag, {data.get('namn', '___')}, överklagar beslutet från {data.get('beslutsmyndighet', '___')} den {data.get('datum', '___')}.").bold = True

    doc.add_heading("Skäl", level=2)
    skäl = f"Beslutet strider mot {data.get('lagrum', '___')}, särskilt i fråga om {data.get('grund', '___')}."
    doc.add_paragraph(skäl)

    doc.add_paragraph("
Ort och datum: __________________________")
    doc.add_paragraph("Underskrift: ____________________________")

    # Format
    for p in doc.paragraphs:
        p.style.font.size = Pt(11)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Spara DOCX
    filepath = f"/mnt/data/{filename}"
    doc.save(filepath)
    return filepath
