"""
doc_strategy_optimizer.py
Genererar strategimärkta Word- och PDF-dokument från juridisk text.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

def skapa_word_dokument(titel, innehåll, strategidata=None, bilagor=None, spara_som="output.docx"):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_heading(titel, 0)

    if strategidata:
        doc.add_paragraph(f"Strategisk profilering – Genererad {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    if strategidata:
        risk = strategidata.get("riskprofil", "Ej angiven")
        prejudikat = strategidata.get("prejudikat", "Ej analyserad")
        doc.add_paragraph(f"🧠 Riskprofil: {risk}")
        doc.add_paragraph(f"⚖️ Prejudikatläge: {prejudikat}")

    doc.add_paragraph("")

    for i, punkt in enumerate(innehåll, 1):
        para = doc.add_paragraph()
        run = para.add_run(f"{i}. {punkt}")
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run.bold = False

    if bilagor:
        doc.add_page_break()
        doc.add_heading("📎 Bilageindex", level=1)
        for i, bil in enumerate(bilagor, 1):
            doc.add_paragraph(f"{i}. {bil.get('namn')} – {bil.get('beskrivning')}")

    doc.save(spara_som)
    return spara_som

def skapa_pdf_via_weasyprint(html_str, spara_som_pdf):
    from weasyprint import HTML
    HTML(string=html_str).write_pdf(spara_som_pdf)
    return spara_som_pdf

def generera_html_version(titel, innehåll, strategidata=None, bilagor=None):
    html = f"<h1>{titel}</h1>"
    if strategidata:
        html += f"<p><strong>🧠 Riskprofil:</strong> {strategidata.get('riskprofil')}<br>"
        html += f"<strong>⚖️ Prejudikatläge:</strong> {strategidata.get('prejudikat')}</p>"

    for i, punkt in enumerate(innehåll, 1):
        html += f"<p><strong>{i}.</strong> {punkt}</p>"

    if bilagor:
        html += "<hr><h2>📎 Bilageindex</h2>"
        for i, bil in enumerate(bilagor, 1):
            html += f"<p>{i}. {bil.get('namn')} – {bil.get('beskrivning')}</p>"

    return html
